#!/usr/bin/env python3
"""Build Word document for MIS assignment - reads content from JSON to avoid quoting issues."""
import json, os, sys

OUTPUT_DIR = r"C:\Users\NewtN\下载"

# Load content
with open(os.path.join(OUTPUT_DIR, 'doc_content.json'), 'r', encoding='utf-8') as f:
    C = json.load(f)

# Now import docx and matplotlib
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch

plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

STUDENT_NAME = C['student_name']
STUDENT_ID = C['student_id']
OUTPUT_FILE = os.path.join(OUTPUT_DIR, f"{STUDENT_ID}-{STUDENT_NAME}-生活后勤系统.docx")


# ============ Diagram generators ============

def fig_org_model(fp):
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis('off')
    ax.set_title('\u7ec4\u7ec7\u89d2\u8272\u5173\u7cfb\u6a21\u578b', fontsize=16, fontweight='bold', pad=15)

    roles = [
        (5, 5.2, '\u62a5\u4fee\u5b66\u751f\n(\u5de5\u5355\u53d1\u8d77\u8005/\u670d\u52a1\u5bf9\u8c61)', '#4A90D9'),
        (1.5, 2.5, '\u540e\u52e4\u53d7\u7406\u5458\n(\u5ba1\u6838\u7533\u8bf7/\u8c03\u5ea6\u534f\u8c03)', '#50B86C'),
        (8.5, 2.5, '\u7ef4\u4fee\u4eba\u5458\n(\u6267\u884c\u7ef4\u4fee/\u786e\u8ba4\u72b6\u6001)', '#E8833A'),
    ]
    for x, y, text, color in roles:
        rect = FancyBboxPatch((x-1.6, y-0.8), 3.2, 1.6, boxstyle="round,pad=0.15",
                              facecolor=color, edgecolor='white', alpha=0.85, linewidth=2)
        ax.add_patch(rect)
        ax.text(x, y, text, ha='center', va='center', fontsize=10, color='white', fontweight='bold')

    arrows = [
        (2.5, 3.2, 3.8, 4.6, '\u63d0\u4ea4\u53d6\u6d88/\u6539\u671f\u7533\u8bf7', '#333', '-'),
        (3.8, 4.3, 2.5, 3.5, '\u5ba1\u6838\u7ed3\u679c\u901a\u77e5', '#666', '--'),
        (7.5, 3.2, 6.2, 4.6, '\u6307\u6d3e/\u901a\u77e5\u7ef4\u4fee\u53d8\u66f4', '#333', '-'),
        (6.2, 4.3, 7.5, 3.5, '\u53cd\u9988\u5de5\u5355\u72b6\u6001', '#666', '--'),
        (7.8, 2.8, 5.8, 2.8, '\u5230\u573a\u7ef4\u4fee', '#999', '-'),
    ]
    for x1, y1, x2, y2, label, color, ls in arrows:
        ax.annotate(label, xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color=color, lw=2, ls=ls),
                    fontsize=9, color=color, ha='center')

    ax.text(5, 0.15, '\u56fe\uff1a\u5de5\u5355\u53d6\u6d88\u4e0e\u6539\u671f\u5b50\u529f\u80fd\u7684\u7ec4\u7ec7\u89d2\u8272\u5173\u7cfb', ha='center', fontsize=10, color='#666', style='italic')
    plt.tight_layout()
    fig.savefig(fp, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()


def fig_er(fp):
    fig, ax = plt.subplots(figsize=(14, 9))
    ax.set_xlim(0, 14); ax.set_ylim(0, 9); ax.axis('off')
    ax.set_title('E-R\u56fe\uff08\u5b9e\u4f53-\u8054\u7cfb\u56fe\uff09', fontsize=16, fontweight='bold', pad=15)

    entities = [
        (2, 7, '\u5b66\u751f\u7528\u6237', ['\u5b66\u53f7(PK)', '\u59d3\u540d', '\u624b\u673a\u53f7', '\u5bbf\u820d\u697c\u680b', '\u5bbf\u820d\u53f7']),
        (7, 7.5, '\u62a5\u4fee\u5de5\u5355', ['\u5de5\u5355\u7f16\u53f7(PK)', '\u5b66\u53f7(FK)', '\u6545\u969c\u9879\u76ee', '\u6545\u969c\u63cf\u8ff0', '\u5de5\u5355\u72b6\u6001', '\u63d0\u4ea4\u65f6\u95f4', '\u9884\u7ea6\u65f6\u95f4']),
        (12, 5.5, '\u9884\u7ea6\u65f6\u6bb5', ['\u65f6\u6bb5\u7f16\u53f7(PK)', '\u65e5\u671f', '\u5f00\u59cb\u65f6\u95f4', '\u7ed3\u675f\u65f6\u95f4', '\u662f\u5426\u53ef\u9009']),
        (2, 3, '\u53d6\u6d88\u539f\u56e0', ['\u539f\u56e0\u7f16\u7801(PK)', '\u539f\u56e0\u63cf\u8ff0', '\u662f\u5426\u9700\u8981\u5ba1\u6838']),
        (7, 2.5, '\u901a\u77e5\u6d88\u606f', ['\u6d88\u606f\u7f16\u53f7(PK)', '\u5de5\u5355\u7f16\u53f7(FK)', '\u5b66\u53f7(FK)', '\u6d88\u606f\u7c7b\u578b', '\u6d88\u606f\u5185\u5bb9', '\u53d1\u9001\u65f6\u95f4', '\u9605\u8bfb\u72b6\u6001']),
    ]
    for x, y, name, attrs in entities:
        h = 0.5 + 0.28 * len(attrs)
        rect = FancyBboxPatch((x-2.0, y-h/2), 4.0, h, boxstyle="round,pad=0.1",
                              facecolor='#EBF5FB', edgecolor='#2E86C1', linewidth=2)
        ax.add_patch(rect)
        ax.text(x, y + h/2 - 0.3, name, ha='center', va='center', fontsize=11, fontweight='bold', color='#1A5276')
        for i, attr in enumerate(attrs):
            ax.text(x, y + h/2 - 0.7 - i*0.28, attr, ha='center', va='center', fontsize=8, color='#333')

    rels = [
        (5, 7.0, '\u63d0\u4ea4', '1', 'n'),
        (9.5, 6.3, '\u5173\u8054\n\u9884\u7ea6', '1', '1'),
        (4.5, 5.0, '\u9009\u62e9\n\u539f\u56e0', 'n', '1'),
        (5, 4.5, '\u89e6\u53d1', '1', 'n'),
    ]
    for x, y, label, c1, c2 in rels:
        diamond = plt.Polygon([(x, y+0.5), (x+1.0, y), (x, y-0.5), (x-1.0, y)],
                              facecolor='#FAD7A1', edgecolor='#E67E22', linewidth=1.5)
        ax.add_patch(diamond)
        ax.text(x, y, label, ha='center', va='center', fontsize=8, fontweight='bold', color='#935116')
        ax.text(x-1.4, y+0.30, c1, fontsize=8, color='#666')
        ax.text(x+1.0, y+0.30, c2, fontsize=8, color='#666')

    ax.text(7, 0.15, '\u56fe\uff1a\u62a5\u4fee\u5de5\u5355\u53d6\u6d88\u4e0e\u6539\u671f\u5b50\u529f\u80fd\u7684E-R\u56fe', ha='center', fontsize=10, color='#666', style='italic')
    plt.tight_layout()
    fig.savefig(fp, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()


def fig_biz_flow(fp):
    fig, ax = plt.subplots(figsize=(16, 10))
    ax.set_xlim(0, 16); ax.set_ylim(0, 10); ax.axis('off')
    ax.set_title('\u4e1a\u52a1\u6d41\u7a0b\u56fe\u2014\u2014\u5de5\u5355\u53d6\u6d88\u4e0e\u9884\u7ea6\u6539\u671f\uff08\u6cf3\u9053\u56fe\uff09', fontsize=14, fontweight='bold', pad=15)

    lanes = ['\u62a5\u4fee\u5b66\u751f', '\u7cfb\u7edf\u5e73\u53f0', '\u540e\u52e4\u53d7\u7406\u5458', '\u7ef4\u4fee\u4eba\u5458']
    lane_colors = ['#D6EAF8', '#D5F5E3', '#FCF3CF', '#FADBD8']
    lane_y = [8.5, 6.5, 4.5, 2.5]
    for i, (lane, y, c) in enumerate(zip(lanes, lane_y, lane_colors)):
        ax.fill_between([0.5, 15.5], y, y+1.5, color=c, alpha=0.4)
        ax.plot([0.5, 15.5], [y, y], color='#999', lw=0.5)
        ax.plot([0.5, 15.5], [y+1.5, y+1.5], color='#999', lw=0.5)
        ax.text(0.8, y+0.75, lane, fontsize=11, fontweight='bold', color='#333', va='center')

    nodes = [
        (2, 9.0, '\u53d1\u8d77\u53d6\u6d88/\n\u6539\u671f\u7533\u8bf7', '#4A90D9'),
        (2, 7.0, '\u6821\u9a8c\u5de5\u5355\u72b6\u6001\n(\u662f\u5426\u5df2\u5f00\u5de5)', '#F39C12'),
        (2, 5.0, '\u5ba1\u6838\u7533\u8bf7\n(\u81ea\u52a8/\u4eba\u5de5)', '#50B86C'),
        (2, 3.0, '\u6536\u5230\u53d8\u66f4\u901a\n\u77e5,\u8c03\u6574\u6392\u73ed', '#E8833A'),
        (6, 9.0, '\u6536\u5230\u5ba1\u6838\n\u7ed3\u679c\u901a\u77e5', '#85C1E9'),
        (10, 7.0, '\u66f4\u65b0\u5de5\u5355\n\u72b6\u6001', '#F8C471'),
        (10, 5.0, '\u5de5\u5355\u72b6\u6001\u540c\u6b65\n+\u8bb0\u5f55\u53d8\u66f4\u65e5\u5fd7', '#82E0AA'),
        (10, 3.0, '\u786e\u8ba4\u65b0\u65f6\u6bb5\n\u6216\u5de5\u5355\u5173\u95ed', '#F0B27A'),
        (13.5, 9.0, '\u7ed3\u675f', '#7FB3D8'),
    ]
    for x, y, text, color in nodes:
        rect = FancyBboxPatch((x-1.5, y-0.55), 3.0, 1.1, boxstyle="round,pad=0.08",
                              facecolor=color, edgecolor='white', alpha=0.9, linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x, y, text, ha='center', va='center', fontsize=8, color='white', fontweight='bold')

    # Decision diamond
    diamond_pts = [(5, 6.75), (5.6, 7.5), (5, 8.25), (4.4, 7.5)]
    diamond = plt.Polygon(diamond_pts, facecolor='#E74C3C', edgecolor='white', alpha=0.85)
    ax.add_patch(diamond)
    ax.text(5, 7.5, '\u5df2\u5f00\u5de5?', ha='center', va='center', fontsize=8, color='white', fontweight='bold')

    arrow_pairs = [
        (3.5, 9.0, 3.5, 8.25), (3.5, 8.25, 4.4, 7.5),
        (5.6, 7.5, 8.5, 9.0), (5, 6.75, 5, 5.55),
        (3.5, 5.0, 5.5, 5.0), (8.5, 7.0, 11.5, 7.0),
        (6, 8.45, 6, 7.55), (11.5, 5.0, 11.5, 3.55),
        (10, 8.45, 10, 7.55), (12, 9.0, 13.5, 9.0),
    ]
    for x1, y1, x2, y2 in arrow_pairs:
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    ax.text(3.8, 8.7, '\u5426(\u672a\u5f00\u5de5\n\u76f4\u63a5\u53d6\u6d88)', fontsize=7, color='#E74C3C', ha='center')
    ax.text(5.3, 6.4, '\u662f(\u8fdb\u5165\u5ba1\u6838)', fontsize=7, color='#E74C3C', ha='center')
    ax.text(8, 0.15, '\u56fe\uff1a\u5de5\u5355\u53d6\u6d88\u4e0e\u6539\u671f\u5b50\u529f\u80fd\u7684\u6cf3\u9053\u5f0f\u4e1a\u52a1\u6d41\u7a0b\u56fe', ha='center', fontsize=10, color='#666', style='italic')
    plt.tight_layout()
    fig.savefig(fp, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()


def fig_dfd0(fp):
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(0, 10); ax.set_ylim(0, 7); ax.axis('off')
    ax.set_title('\u9876\u5c42\u6570\u636e\u6d41\u7a0b\u56fe\uff08DFD-0\uff09', fontsize=16, fontweight='bold', pad=15)

    for x, y, text in [(1.5, 3.5, '\u62a5\u4fee\u5b66\u751f'), (8.5, 3.5, '\u540e\u52e4\u53d7\u7406\u5458')]:
        rect = FancyBboxPatch((x-1.2, y-0.45), 2.4, 0.9, boxstyle="round,pad=0.08",
                              facecolor='#D5F5E3', edgecolor='#27AE60', linewidth=2)
        ax.add_patch(rect)
        ax.text(x, y, text, ha='center', va='center', fontsize=11, fontweight='bold')

    process = FancyBboxPatch((3.5, 2.5), 3.0, 2.0, boxstyle="round,pad=0.2",
                              facecolor='#D6EAF8', edgecolor='#2E86C1', linewidth=2)
    ax.add_patch(process)
    ax.text(5, 3.5, '\u5de5\u5355\u53d6\u6d88\u4e0e\n\u6539\u671f\u5904\u7406\n\u7cfb\u7edf', ha='center', va='center', fontsize=12, fontweight='bold')

    flows = [
        (2.7, 3.8, 3.5, 3.8, '\u53d6\u6d88/\u6539\u671f\u7533\u8bf7'),
        (6.5, 3.2, 7.3, 3.2, '\u5ba1\u6838\u7ed3\u679c\u901a\u77e5'),
        (2.7, 3.2, 3.5, 3.2, '\u5de5\u5355\u72b6\u6001\u67e5\u8be2'),
        (6.5, 3.8, 7.3, 3.8, '\u53d8\u66f4\u901a\u77e5'),
    ]
    for x1, y1, x2, y2, label in flows:
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1), arrowprops=dict(arrowstyle='->', color='#555', lw=2))
        ax.text((x1+x2)/2, (y1+y2)/2 + 0.2, label, fontsize=9, color='#333', ha='center')

    ax.text(5, 1.5, '\u56fe\uff1a\u9876\u5c42\u6570\u636e\u6d41\u7a0b\u56fe', ha='center', fontsize=10, color='#666', style='italic')
    plt.tight_layout()
    fig.savefig(fp, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()


def fig_dfd1(fp):
    fig, ax = plt.subplots(figsize=(12, 8))
    ax.set_xlim(0, 12); ax.set_ylim(0, 8); ax.axis('off')
    ax.set_title('\u4e00\u5c42\u6570\u636e\u6d41\u7a0b\u56fe\uff08DFD-1\uff09', fontsize=16, fontweight='bold', pad=15)

    for x, y, text in [(1.5, 4, '\u62a5\u4fee\u5b66\u751f'), (10.5, 4, '\u540e\u52e4\u53d7\u7406\u5458')]:
        rect = FancyBboxPatch((x-1.0, y-0.4), 2.0, 0.8, boxstyle="round,pad=0.08",
                              facecolor='#D5F5E3', edgecolor='#27AE60', linewidth=2)
        ax.add_patch(rect)
        ax.text(x, y, text, ha='center', va='center', fontsize=10, fontweight='bold')

    processes = [
        (4, 6, 'P1\n\u7533\u8bf7\u63d0\u4ea4\n\u5904\u7406'),
        (6, 4, 'P2\n\u5ba1\u6838\u5904\u7406'),
        (8, 6, 'P3\n\u901a\u77e5\u63a8\u9001\n\u5904\u7406'),
    ]
    for x, y, text in processes:
        circle = plt.Circle((x, y), 0.9, facecolor='#D6EAF8', edgecolor='#2E86C1', linewidth=2)
        ax.add_patch(circle)
        ax.text(x, y, text, ha='center', va='center', fontsize=9, fontweight='bold')

    stores = [
        (4, 2, 'D1 \u5de5\u5355\u6570\u636e'),
        (6, 2, 'D2 \u53d6\u6d88\u539f\u56e0\n    \u5b57\u5178'),
        (8, 2, 'D3 \u901a\u77e5\u6d88\u606f\n    \u8bb0\u5f55'),
    ]
    for x, y, text in stores:
        rect = FancyBboxPatch((x-1.1, y-0.45), 2.2, 0.9, boxstyle="round,pad=0.05",
                              facecolor='#FCF3CF', edgecolor='#F1C40F', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x, y, text, ha='center', va='center', fontsize=8, color='#7D6608')

    flow_pairs = [
        (2.5, 4.2, 3.1, 5.5), (4.9, 5.1, 5.1, 4.9), (6.9, 4.9, 7.1, 5.1),
        (8.9, 5.5, 9.5, 4.2), (4, 4.8, 4, 2.9), (6, 4.8, 6, 2.9), (8, 4.8, 8, 2.9),
    ]
    for x1, y1, x2, y2 in flow_pairs:
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1), arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    ax.text(2.8, 4.8, '\u7533\u8bf7\u6570\u636e', fontsize=8, color='#333')
    ax.text(5, 5.3, '\u6821\u9a8c\u540e\n\u5de5\u5355\u6570\u636e', fontsize=7, color='#333', ha='center')
    ax.text(7, 5.3, '\u5ba1\u6838\u7ed3\u679c', fontsize=8, color='#333', ha='center')
    ax.text(9.2, 4.8, '\u901a\u77e5\u6d88\u606f', fontsize=8, color='#333')
    ax.text(6, 1.0, '\u56fe\uff1a\u4e00\u5c42\u6570\u636e\u6d41\u7a0b\u56fe', ha='center', fontsize=10, color='#666', style='italic')
    plt.tight_layout()
    fig.savefig(fp, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()


def fig_func_struct(fp):
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.set_xlim(0, 14); ax.set_ylim(0, 8); ax.axis('off')
    ax.set_title('\u529f\u80fd\u7ed3\u6784\u56fe\u2014\u2014\u5de5\u5355\u53d6\u6d88\u4e0e\u6539\u671f\u5b50\u7cfb\u7edf', fontsize=16, fontweight='bold', pad=15)

    root = (7, 7.3)
    rect = FancyBboxPatch((root[0]-2.2, root[1]-0.35), 4.4, 0.7, boxstyle="round,pad=0.1",
                          facecolor='#2E86C1', edgecolor='white', linewidth=2)
    ax.add_patch(rect)
    ax.text(root[0], root[1], '\u5de5\u5355\u53d6\u6d88\u4e0e\u6539\u671f\u5b50\u7cfb\u7edf', fontsize=12, color='white', fontweight='bold', ha='center', va='center')

    level1 = [
        (3, 5.8, '\u5b66\u751f\u7aef\n\u5de5\u5355\u64cd\u4f5c\u6a21\u5757'),
        (7, 5.8, '\u7ba1\u7406\u7aef\n\u5ba1\u6838\u5904\u7406\u6a21\u5757'),
        (11, 5.8, '\u6d88\u606f\u901a\u77e5\n\u6a21\u5757'),
    ]
    for x, y, text in level1:
        rect = FancyBboxPatch((x-1.5, y-0.5), 3.0, 1.0, boxstyle="round,pad=0.1",
                              facecolor='#3498DB', edgecolor='white', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x, y, text, fontsize=10, color='white', fontweight='bold', ha='center', va='center')
        ax.plot([root[0], x], [root[1]-0.35, y+0.5], color='#555', lw=1.5)

    # Level 2
    l2data = [
        (3, [(0.8, 4.0, '\u5de5\u5355\u53d6\u6d88\n\u7533\u8bf7'), (3.0, 4.0, '\u9884\u7ea6\u6539\u671f\n\u7533\u8bf7'), (5.2, 4.0, '\u7533\u8bf7\u8fdb\u5ea6\n\u67e5\u8be2')]),
        (7, [(4.8, 4.0, '\u53d6\u6d88\u539f\u56e0\n\u5ba1\u6838'), (7.0, 4.0, '\u6539\u671f\u65f6\u6bb5\n\u5ba1\u6279'), (9.2, 4.0, '\u5de5\u5355\u72b6\u6001\n\u7ba1\u7406')]),
        (11, [(9.8, 4.0, '\u4f01\u4e1a\u5fae\u4fe1\n\u6d88\u606f\u63a8\u9001'), (11.0, 4.0, '\u7ad9\u5185\u901a\u77e5\n\u7ba1\u7406'), (12.2, 4.0, '\u6d88\u606f\u6a21\u677f\n\u914d\u7f6e')]),
    ]
    for px, children in l2data:
        for cx, cy, text in children:
            rect = FancyBboxPatch((cx-0.7, cy-0.45), 1.4, 0.9, boxstyle="round,pad=0.05",
                                  facecolor='#85C1E9', edgecolor='white', linewidth=1)
            ax.add_patch(rect)
            ax.text(cx, cy, text, fontsize=8, color='#1A5276', ha='center', va='center', fontweight='bold')
            ax.plot([px, cx], [5.3, cy+0.45], color='#999', lw=1)

    # Level 3
    l3data = [
        (0.8, [(0.15, 2.4, '\u9009\u62e9\u53d6\n\u6d88\u539f\u56e0'), (0.8, 2.4, '\u786e\u8ba4\u53d6\n\u6d88\u64cd\u4f5c'), (1.45, 2.4, '\u67e5\u770b\u53d6\n\u6d88\u7ed3\u679c')]),
        (3.0, [(2.35, 2.4, '\u9009\u62e9\u65b0\n\u65f6\u95f4\u6bb5'), (3.0, 2.4, '\u586b\u5199\u6539\n\u671f\u539f\u56e0'), (3.65, 2.4, '\u63d0\u4ea4\u6539\n\u671f\u7533\u8bf7')]),
        (5.2, [(4.55, 2.4, '\u53d6\u6d88\u7533\n\u8bf7\u8fdb\u5ea6'), (5.2, 2.4, '\u6539\u671f\u7533\n\u8bf7\u8fdb\u5ea6'), (5.85, 2.4, '\u5386\u53f2\u7533\n\u8bf7\u8bb0\u5f55')]),
        (4.8, [(4.15, 2.4, '\u81ea\u52a8\u5ba1\n\u6838\u89c4\u5219'), (4.8, 2.4, '\u4eba\u5de5\u5ba1\n\u6838\u754c\u9762'), (5.45, 2.4, '\u6279\u91cf\u5904\n\u7406\u64cd\u4f5c')]),
        (7.0, [(6.35, 2.4, '\u65f6\u6bb5\u53ef\n\u7528\u6027\u7ef4\u62a4'), (7.0, 2.4, '\u6539\u671f\u5ba1\n\u6279\u64cd\u4f5c'), (7.65, 2.4, '\u51b2\u7a81\u68c0\n\u6d4b\u5904\u7406')]),
        (9.2, [(8.35, 2.4, '\u72b6\u6001\u53d8\n\u66f4\u8bb0\u5f55'), (9.2, 2.4, '\u7ef4\u4fee\u4eba\n\u5458\u901a\u77e5'), (10.05,2.4, '\u5de5\u5355\u5173\n\u95ed\u5f52\u6863')]),
        (9.8, [(9.15, 2.4, '\u53d6\u6d88\u6210\n\u529f\u901a\u77e5'), (9.8, 2.4, '\u6539\u671f\u786e\n\u8ba4\u901a\u77e5'), (10.45,2.4, '\u50ac\u529e\u63d0\n\u9192\u901a\u77e5')]),
        (11.0,[(10.35,2.4, '\u6d88\u606f\u5217\n\u8868\u7ba1\u7406'), (11.0,2.4, '\u5df2\u8bfb/\u672a\n\u8bfb\u8ffd\u8e2a'), (11.65,2.4, '\u6d88\u606f\u641c\n\u7d22\u8fc7\u6ee4')]),
        (12.2,[(11.55,2.4, '\u901a\u77e5\u6a21\n\u677f\u7f16\u8f91'), (12.2,2.4, '\u53d8\u91cf\u66ff\n\u6362\u89c4\u5219'), (12.85,2.4, '\u53d1\u9001\u65f6\n\u673a\u914d\u7f6e')]),
    ]
    for px, children in l3data:
        for cx, cy, text in children:
            rect = FancyBboxPatch((cx-0.32, cy-0.38), 0.64, 0.76, boxstyle="round,pad=0.03",
                                  facecolor='#D6EAF8', edgecolor='#AED6F1', linewidth=0.8)
            ax.add_patch(rect)
            ax.text(cx, cy, text, fontsize=6.5, color='#1A5276', ha='center', va='center')
            ax.plot([px, cx], [3.55, cy+0.38], color='#CCC', lw=0.7)

    ax.text(7, 0.15, '\u56fe\uff1a\u5de5\u5355\u53d6\u6d88\u4e0e\u6539\u671f\u5b50\u7cfb\u7edf\u7684\u4e09\u5c42\u529f\u80fd\u7ed3\u6784\u56fe', ha='center', fontsize=10, color='#666', style='italic')
    plt.tight_layout()
    fig.savefig(fp, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()


# ============ Word builder helpers ============

def add_styled_table(doc, headers, rows, col_widths=None, header_color='1F4E79'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), header_color)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if i % 2 == 1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'EBF5FB')
                shading.set(qn('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'SimHei'
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), 'SimHei')
        rPr.insert(0, rFonts)
    return heading


def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'SimSun'
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'SimSun')
    rPr.insert(0, rFonts)
    if bold:
        run.font.bold = True
    return p


def add_image(doc, filepath, width_cm=14):
    if os.path.exists(filepath):
        doc.add_picture(filepath, width=Cm(width_cm))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ============ Main document builder ============

def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ---- Cover ----
    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(C['cover']['title'])
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.name = 'SimHei'
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run(C['cover']['subtitle'])
    run.font.size = Pt(16)
    run.font.name = 'KaiTi'
    run.font.color.rgb = RGBColor(51, 51, 51)

    doc.add_paragraph(); doc.add_paragraph()
    for item in C['cover']['info']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        run.font.size = Pt(13)
        run.font.name = 'SimSun'

    doc.add_page_break()

    # ---- TOC placeholder ----
    add_heading_styled(doc, '\u76ee\u5f55', level=1)
    add_para(doc, '\uff08\u8bf7\u5728Word\u4e2d\u63d2\u5165\u81ea\u52a8\u76ee\u5f55\uff1a\u5f15\u7528 \u2192 \u76ee\u5f55 \u2192 \u81ea\u52a8\u76ee\u5f55\uff09')
    doc.add_page_break()

    # ---- Generate diagrams first ----
    img_dir = OUTPUT_DIR
    org_img = os.path.join(img_dir, '_fig_org.png')
    er_img = os.path.join(img_dir, '_fig_er.png')
    biz_img = os.path.join(img_dir, '_fig_biz.png')
    dfd0_img = os.path.join(img_dir, '_fig_dfd0.png')
    dfd1_img = os.path.join(img_dir, '_fig_dfd1.png')
    func_img = os.path.join(img_dir, '_fig_func.png')

    fig_org_model(org_img)
    fig_er(er_img)
    fig_biz_flow(biz_img)
    fig_dfd0(dfd0_img)
    fig_dfd1(dfd1_img)
    fig_func_struct(func_img)

    # ====== Part 1 ======
    add_heading_styled(doc, C['part1']['title'], level=1)
    for sec in C['part1']['sections']:
        add_heading_styled(doc, sec['heading'], level=2)
        for para in sec['content']:
            if isinstance(para, str):
                add_para(doc, para, indent=sec.get('indent', True))
            elif isinstance(para, dict):
                if para.get('type') == 'bold':
                    add_para(doc, para['text'], bold=True)
                elif para.get('type') == 'table':
                    add_styled_table(doc, para['headers'], para['rows'], para.get('col_widths'))
                    if para.get('caption'):
                        add_para(doc, para['caption'], indent=True)
    doc.add_page_break()

    # ====== Part 2 ======
    add_heading_styled(doc, C['part2']['title'], level=1)
    for sec in C['part2']['sections']:
        add_heading_styled(doc, sec['heading'], level=2)
        for para in sec['content']:
            if isinstance(para, str):
                add_para(doc, para, indent=sec.get('indent', True))
            elif isinstance(para, dict):
                t = para.get('type', '')
                if t == 'bold':
                    add_para(doc, para['text'], bold=True)
                elif t == 'table':
                    add_styled_table(doc, para['headers'], para['rows'], para.get('col_widths'))
                    if para.get('caption'):
                        add_para(doc, para['caption'], indent=True)
                elif t == 'image':
                    img_map = {'org': org_img, 'er': er_img}
                    add_image(doc, img_map.get(para['id'], ''))
    doc.add_page_break()

    # ====== Part 3 ======
    add_heading_styled(doc, C['part3']['title'], level=1)
    for sec in C['part3']['sections']:
        add_heading_styled(doc, sec['heading'], level=2)
        for para in sec['content']:
            if isinstance(para, str):
                add_para(doc, para, indent=sec.get('indent', True))
            elif isinstance(para, dict):
                t = para.get('type', '')
                if t == 'bold':
                    add_para(doc, para['text'], bold=True)
                elif t == 'table':
                    add_styled_table(doc, para['headers'], para['rows'], para.get('col_widths'))
                    if para.get('caption'):
                        add_para(doc, para['caption'], indent=True)
                elif t == 'image':
                    img_map = {'biz': biz_img, 'dfd0': dfd0_img, 'dfd1': dfd1_img}
                    add_image(doc, img_map.get(para['id'], ''))
    doc.add_page_break()

    # ====== Part 4 ======
    add_heading_styled(doc, C['part4']['title'], level=1)
    for sec in C['part4']['sections']:
        add_heading_styled(doc, sec['heading'], level=2)
        for para in sec['content']:
            if isinstance(para, str):
                add_para(doc, para, indent=sec.get('indent', True))
            elif isinstance(para, dict):
                t = para.get('type', '')
                if t == 'bold':
                    add_para(doc, para['text'], bold=True)
                elif t == 'table':
                    add_styled_table(doc, para['headers'], para['rows'], para.get('col_widths'))
                    if para.get('caption'):
                        add_para(doc, para['caption'], indent=True)
                elif t == 'image':
                    add_image(doc, func_img)
    doc.add_page_break()

    # ====== Part 5: Conclusion ======
    add_heading_styled(doc, C['conclusion']['title'], level=1)
    for sec in C['conclusion']['sections']:
        if isinstance(sec, str):
            add_para(doc, sec, indent=True)
        elif isinstance(sec, dict) and sec.get('type') == 'bold':
            add_para(doc, sec['text'], bold=True)
    for p in C['conclusion']['paragraphs']:
        add_para(doc, p, indent=True)
    doc.add_page_break()

    # ====== Appendix ======
    add_heading_styled(doc, C['appendix']['title'], level=1)
    for item in C['appendix']['items']:
        add_para(doc, item)

    # ---- Save ----
    doc.save(OUTPUT_FILE)
    print(f'Document saved: {OUTPUT_FILE}')
    return OUTPUT_FILE


if __name__ == '__main__':
    build()
