#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成管理信息系统大作业Word文档 — 华中科技大学生活后勤系统升级改造"""

import os, sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Arc, ConnectionPatch
import numpy as np

# ============ 配置 ============
plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

OUTPUT_DIR = r"C:\Users\NewtN\下载"
STUDENT_NAME = "鞠京芮"
STUDENT_ID = "________"  # 请填写学号
OUTPUT_FILE = os.path.join(OUTPUT_DIR, f"{STUDENT_ID}-{STUDENT_NAME}-生活后勤系统.docx")

# ============ 1. 图像生成函数 ============

def fig_organization_model(filepath):
    """组织模型——角色关系图"""
    fig, ax = plt.subplots(1, 1, figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis('off')
    ax.set_title("组织角色关系模型", fontsize=16, fontweight='bold', pad=15)

    # 三个角色框
    roles = [
        (5, 5.2, "报修学生\n(工单发起者/服务对象)", "#4A90D9"),
        (1.5, 2.5, "后勤受理员\n(审核申请/调度协调)", "#50B86C"),
        (8.5, 2.5, "维修人员\n(执行维修/确认状态)", "#E8833A"),
    ]
    for x, y, text, color in roles:
        rect = FancyBboxPatch((x-1.6, y-0.8), 3.2, 1.6, boxstyle="round,pad=0.15",
                              facecolor=color, edgecolor='white', alpha=0.85, linewidth=2)
        ax.add_patch(rect)
        ax.text(x, y, text, ha='center', va='center', fontsize=10, color='white', fontweight='bold')

    # 连线标注
    ax.annotate("提交取消/改期申请", xy=(2.5, 3.2), xytext=(3.8, 4.6),
                arrowprops=dict(arrowstyle='->', color='#333', lw=2), fontsize=9, color='#333', ha='center')
    ax.annotate("审核结果通知", xy=(3.8, 4.3), xytext=(2.5, 3.5),
                arrowprops=dict(arrowstyle='->', color='#333', lw=2, ls='--'), fontsize=9, color='#666', ha='center')
    ax.annotate("指派/通知维修变更", xy=(7.5, 3.2), xytext=(6.2, 4.6),
                arrowprops=dict(arrowstyle='->', color='#333', lw=2), fontsize=9, color='#333', ha='center')
    ax.annotate("反馈工单状态", xy=(6.2, 4.3), xytext=(7.5, 3.5),
                arrowprops=dict(arrowstyle='->', color='#333', lw=2, ls='--'), fontsize=9, color='#666', ha='center')
    ax.annotate("到场维修", xy=(7.8, 2.8), xytext=(5.8, 2.8),
                arrowprops=dict(arrowstyle='<->', color='#999', lw=1.5), fontsize=9, color='#999', ha='center')

    ax.text(5, 0.15, "图：工单取消与改期子功能的组织角色关系", ha='center', fontsize=10, color='#666', style='italic')
    plt.tight_layout()
    fig.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()

def fig_er_diagram(filepath):
    """E-R图"""
    fig, ax = plt.subplots(1, 1, figsize=(14, 9))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.set_title("实体-联系图（E-R图）", fontsize=16, fontweight='bold', pad=15)

    entities = [
        (2, 7, "学生用户", ["学号(PK)", "姓名", "手机号", "宿舍楼栋", "宿舍号"]),
        (7, 7.5, "报修工单", ["工单编号(PK)", "学号(FK)", "故障项目", "故障描述", "工单状态", "提交时间", "预约时间"]),
        (12, 5.5, "预约时段", ["时段编号(PK)", "日期", "开始时间", "结束时间", "是否可选"]),
        (2, 3, "取消原因", ["原因编码(PK)", "原因描述", "是否需要审核"]),
        (7, 2.5, "通知消息", ["消息编号(PK)", "工单编号(FK)", "学号(FK)", "消息类型", "消息内容", "发送时间", "阅读状态"]),
    ]

    for x, y, name, attrs in entities:
        h = 0.5 + 0.28 * len(attrs)
        rect = FancyBboxPatch((x-2.0, y-h/2), 4.0, h, boxstyle="round,pad=0.1",
                              facecolor='#EBF5FB', edgecolor='#2E86C1', linewidth=2)
        ax.add_patch(rect)
        ax.text(x, y + h/2 - 0.3, name, ha='center', va='center', fontsize=11, fontweight='bold', color='#1A5276')
        for i, attr in enumerate(attrs):
            ax.text(x, y + h/2 - 0.7 - i*0.28, attr, ha='center', va='center', fontsize=8, color='#333')

    # 菱形关系
    relationships = [
        (5, 7.0, "提交", "1", "n"),
        (9.5, 6.3, "关联\n预约", "1", "1"),
        (4.5, 5.0, "选择\n原因", "n", "1"),
        (5, 4.5, "触发", "1", "n"),
    ]
    for x, y, label, c1, c2 in relationships:
        diamond = plt.Polygon([(x, y+0.5), (x+1.0, y), (x, y-0.5), (x-1.0, y)],
                              facecolor='#FAD7A1', edgecolor='#E67E22', linewidth=1.5)
        ax.add_patch(diamond)
        ax.text(x, y, label, ha='center', va='center', fontsize=8, fontweight='bold', color='#935116')
        ax.text(x-1.4, y+0.30, c1, fontsize=8, color='#666')
        ax.text(x+1.0, y+0.30, c2, fontsize=8, color='#666')

    ax.text(7, 0.15, "图：报修工单取消与改期子功能的E-R图", ha='center', fontsize=10, color='#666', style='italic')
    plt.tight_layout()
    fig.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()


def fig_business_flow(filepath):
    """泳道式业务流程图"""
    fig, ax = plt.subplots(1, 1, figsize=(16, 10))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 10)
    ax.axis('off')
    ax.set_title("业务流程图——工单取消与预约改期（泳道图）", fontsize=14, fontweight='bold', pad=15)

    # 泳道
    lanes = ["报修学生", "系统平台", "后勤受理员", "维修人员"]
    lane_colors = ["#D6EAF8", "#D5F5E3", "#FCF3CF", "#FADBD8"]
    lane_y = [8.5, 6.5, 4.5, 2.5]
    for i, (lane, y, c) in enumerate(zip(lanes, lane_y, lane_colors)):
        ax.fill_between([0.5, 15.5], y, y+1.5, color=c, alpha=0.4)
        ax.plot([0.5, 15.5], [y, y], color='#999', lw=0.5)
        ax.plot([0.5, 15.5], [y+1.5, y+1.5], color='#999', lw=0.5)
        ax.text(0.8, y+0.75, lane, fontsize=11, fontweight='bold', color='#333', va='center')

    # 流程节点（简化为框图）
    nodes = [
        (2, 9.0, "发起取消/\n改期申请", "#4A90D9"),
        (2, 7.0, "校验工单状态\n(是否已开工)", "#F39C12"),
        (2, 5.0, "审核申请\n(自动/人工)", "#50B86C"),
        (2, 3.0, "收到变更通\n知,调整排班", "#E8833A"),
        (6, 9.0, "收到审核\n结果通知", "#85C1E9"),
        (10, 7.0, "更新工单\n状态", "#F8C471"),
        (10, 5.0, "工单状态同步\n+记录变更日志", "#82E0AA"),
        (10, 3.0, "确认新时段\n或工单关闭", "#F0B27A"),
        (13.5, 9.0, "结束", "#7FB3D8"),
    ]

    for x, y, text, color in nodes:
        rect = FancyBboxPatch((x-1.5, y-0.55), 3.0, 1.1, boxstyle="round,pad=0.08",
                              facecolor=color, edgecolor='white', alpha=0.9, linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x, y, text, ha='center', va='center', fontsize=8, color='white', fontweight='bold')

    # 判定节点
    diamond_pts = [(5, 6.75), (5.6, 7.5), (5, 8.25), (4.4, 7.5)]
    diamond = plt.Polygon(diamond_pts, facecolor='#E74C3C', edgecolor='white', alpha=0.85)
    ax.add_patch(diamond)
    ax.text(5, 7.5, "已开工?", ha='center', va='center', fontsize=8, color='white', fontweight='bold')

    # 箭头
    arrows = [
        (3.5, 9.0, 3.5, 8.25), (3.5, 8.25, 4.4, 7.5),
        (5.6, 7.5, 8.5, 9.0),  # 否→直接取消
        (5, 6.75, 5, 5.55),     # 是→审核
        (3.5, 5.0, 5.5, 5.0),   # 审核完→状态更新
        (8.5, 7.0, 11.5, 7.0),
        (6, 8.45, 6, 7.55),     # 更新状态后通知
        (11.5, 5.0, 11.5, 3.55),
        (10, 8.45, 10, 7.55),
        (12, 9.0, 13.5, 9.0),
    ]
    for x1, y1, x2, y2 in arrows:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    # 分支标签
    ax.text(3.8, 8.7, "否(未开工\n直接取消)", fontsize=7, color='#E74C3C', ha='center')
    ax.text(5.3, 6.4, "是(进入审核)", fontsize=7, color='#E74C3C', ha='center')

    ax.text(8, 0.15, "图：工单取消与改期子功能的泳道式业务流程图", ha='center', fontsize=10, color='#666', style='italic')
    plt.tight_layout()
    fig.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()


def fig_data_flow_top(filepath):
    """顶层数据流程图"""
    fig, ax = plt.subplots(1, 1, figsize=(10, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis('off')
    ax.set_title("顶层数据流程图（DFD-0）", fontsize=16, fontweight='bold', pad=15)

    # 外部实体（矩形）
    ext_entities = [
        (1.5, 3.5, "报修学生"),
        (8.5, 3.5, "后勤受理员"),
    ]
    for x, y, text in ext_entities:
        rect = FancyBboxPatch((x-1.2, y-0.45), 2.4, 0.9, boxstyle="round,pad=0.08",
                              facecolor='#D5F5E3', edgecolor='#27AE60', linewidth=2)
        ax.add_patch(rect)
        ax.text(x, y, text, ha='center', va='center', fontsize=11, fontweight='bold')

    # 核心处理（圆角矩形）
    process = FancyBboxPatch((3.5, 2.5), 3.0, 2.0, boxstyle="round,pad=0.2",
                              facecolor='#D6EAF8', edgecolor='#2E86C1', linewidth=2)
    ax.add_patch(process)
    ax.text(5, 3.5, "工单取消与\n改期处理\n系统", ha='center', va='center', fontsize=12, fontweight='bold')

    # 数据流箭头
    flows = [
        (2.7, 3.8, 3.5, 3.8, "取消/改期申请"),
        (6.5, 3.2, 7.3, 3.2, "审核结果通知"),
        (2.7, 3.2, 3.5, 3.2, "工单状态查询"),
        (6.5, 3.8, 7.3, 3.8, "变更通知"),
    ]
    for x1, y1, x2, y2, label in flows:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=2))
        ax.text((x1+x2)/2, (y1+y2)/2 + 0.2, label, fontsize=9, color='#333', ha='center')

    ax.text(5, 1.5, "图：顶层数据流程图", ha='center', fontsize=10, color='#666', style='italic')
    plt.tight_layout()
    fig.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()


def fig_data_flow_level1(filepath):
    """一层数据流程图"""
    fig, ax = plt.subplots(1, 1, figsize=(12, 8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_title("一层数据流程图（DFD-1）", fontsize=16, fontweight='bold', pad=15)

    # 外部实体
    ext_entities = [
        (1.5, 4, "报修学生"),
        (10.5, 4, "后勤受理员"),
    ]
    for x, y, text in ext_entities:
        rect = FancyBboxPatch((x-1.0, y-0.4), 2.0, 0.8, boxstyle="round,pad=0.08",
                              facecolor='#D5F5E3', edgecolor='#27AE60', linewidth=2)
        ax.add_patch(rect)
        ax.text(x, y, text, ha='center', va='center', fontsize=10, fontweight='bold')

    # 三个处理过程
    processes = [
        (4, 6, "P1\n申请提交\n处理"),
        (6, 4, "P2\n审核处理"),
        (8, 6, "P3\n通知推送\n处理"),
    ]
    for x, y, text in processes:
        circle = plt.Circle((x, y), 0.9, facecolor='#D6EAF8', edgecolor='#2E86C1', linewidth=2)
        ax.add_patch(circle)
        ax.text(x, y, text, ha='center', va='center', fontsize=9, fontweight='bold')

    # 数据存储（开口矩形）
    stores = [
        (4, 2, "D1 工单数据"),
        (6, 2, "D2 取消原因\n    字典"),
        (8, 2, "D3 通知消息\n    记录"),
    ]
    for x, y, text in stores:
        rect = FancyBboxPatch((x-1.1, y-0.45), 2.2, 0.9, boxstyle="round,pad=0.05",
                              facecolor='#FCF3CF', edgecolor='#F1C40F', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x, y, text, ha='center', va='center', fontsize=8, color='#7D6608')

    # 箭头
    flow_arrows = [
        (2.5, 4.2, 3.1, 5.5),   # 学生→P1
        (4.9, 5.1, 5.1, 4.9),   # P1→P2
        (6.9, 4.9, 7.1, 5.1),   # P2→P3
        (8.9, 5.5, 9.5, 4.2),   # P3→受理员
        (4, 4.8, 4, 2.9),       # P2↔D1
        (6, 4.8, 6, 2.9),       # P2↔D2
        (8, 4.8, 8, 2.9),       # P3↔D3
    ]
    for x1, y1, x2, y2 in flow_arrows:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    ax.text(2.8, 4.8, "申请数据", fontsize=8, color='#333')
    ax.text(5, 5.3, "校验后\n工单数据", fontsize=7, color='#333', ha='center')
    ax.text(7, 5.3, "审核结果", fontsize=8, color='#333', ha='center')
    ax.text(9.2, 4.8, "通知消息", fontsize=8, color='#333')

    ax.text(6, 1.0, "图：一层数据流程图", ha='center', fontsize=10, color='#666', style='italic')
    plt.tight_layout()
    fig.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()


def fig_function_structure(filepath):
    """功能结构图（树状图）"""
    fig, ax = plt.subplots(1, 1, figsize=(14, 8))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_title("功能结构图——工单取消与改期子系统", fontsize=16, fontweight='bold', pad=15)

    # 根节点
    root = (7, 7.3)
    rect = FancyBboxPatch((root[0]-2.2, root[1]-0.35), 4.4, 0.7, boxstyle="round,pad=0.1",
                          facecolor='#2E86C1', edgecolor='white', linewidth=2)
    ax.add_patch(rect)
    ax.text(root[0], root[1], "工单取消与改期子系统", fontsize=12, color='white', fontweight='bold', ha='center', va='center')

    # 一级模块
    level1 = [
        (3, 5.8, "学生端\n工单操作模块"),
        (7, 5.8, "管理端\n审核处理模块"),
        (11, 5.8, "消息通知\n模块"),
    ]
    for x, y, text in level1:
        rect = FancyBboxPatch((x-1.5, y-0.5), 3.0, 1.0, boxstyle="round,pad=0.1",
                              facecolor='#3498DB', edgecolor='white', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x, y, text, fontsize=10, color='white', fontweight='bold', ha='center', va='center')
        ax.plot([root[0], x], [root[1]-0.35, y+0.5], color='#555', lw=1.5)

    # 二级功能
    level2_data = [
        (3, [(0.8, 4.0, "工单取消\n申请"), (3.0, 4.0, "预约改期\n申请"), (5.2, 4.0, "申请进度\n查询")]),
        (7, [(4.8, 4.0, "取消原因\n审核"), (7.0, 4.0, "改期时段\n审批"), (9.2, 4.0, "工单状态\n管理")]),
        (11, [(9.8, 4.0, "企业微信\n消息推送"), (11.0, 4.0, "站内通知\n管理"), (12.2, 4.0, "消息模板\n配置")]),
    ]
    for parent_x, children in level2_data:
        for cx, cy, text in children:
            rect = FancyBboxPatch((cx-0.7, cy-0.45), 1.4, 0.9, boxstyle="round,pad=0.05",
                                  facecolor='#85C1E9', edgecolor='white', linewidth=1)
            ax.add_patch(rect)
            ax.text(cx, cy, text, fontsize=8, color='#1A5276', ha='center', va='center', fontweight='bold')
            ax.plot([parent_x, cx], [5.3, cy+0.45], color='#999', lw=1)

    # 三级功能
    level3_data = [
        (0.8, [(0.15, 2.4, "选择取\n消原因"), (0.8, 2.4, "确认取\n消操作"), (1.45, 2.4, "查看取\n消结果")]),
        (3.0, [(2.35, 2.4, "选择新\n时间段"), (3.0, 2.4, "填写改\n期原因"), (3.65, 2.4, "提交改\n期申请")]),
        (5.2, [(4.55, 2.4, "取消申\n请进度"), (5.2, 2.4, "改期申\n请进度"), (5.85, 2.4, "历史申\n请记录")]),
        (4.8, [(4.15, 2.4, "自动审\n核规则"), (4.8, 2.4, "人工审\n核界面"), (5.45, 2.4, "批量处\n理操作")]),
        (7.0, [(6.35, 2.4, "时段可\n用性维护"), (7.0, 2.4, "改期审\n批操作"), (7.65, 2.4, "冲突检\n测处理")]),
        (9.2, [(8.35, 2.4, "状态变\n更记录"), (9.2, 2.4, "维修人\n员通知"), (10.05,2.4, "工单关\n闭归档")]),
        (9.8, [(9.15, 2.4, "取消成\n功通知"), (9.8, 2.4, "改期确\n认通知"), (10.45,2.4, "催办提\n醒通知")]),
        (11.0,[(10.35,2.4, "消息列\n表管理"), (11.0,2.4, "已读/未\n读追踪"), (11.65,2.4, "消息搜\n索过滤")]),
        (12.2,[(11.55,2.4, "通知模\n板编辑"), (12.2,2.4, "变量替\n换规则"), (12.85,2.4, "发送时\n机配置")]),
    ]
    for parent_x, children in level3_data:
        for cx, cy, text in children:
            rect = FancyBboxPatch((cx-0.32, cy-0.38), 0.64, 0.76, boxstyle="round,pad=0.03",
                                  facecolor='#D6EAF8', edgecolor='#AED6F1', linewidth=0.8)
            ax.add_patch(rect)
            ax.text(cx, cy, text, fontsize=6.5, color='#1A5276', ha='center', va='center')
            ax.plot([parent_x, cx], [3.55, cy+0.38], color='#CCC', lw=0.7)

    ax.text(7, 0.15, "图：工单取消与改期子系统的三层功能结构图", ha='center', fontsize=10, color='#666', style='italic')
    plt.tight_layout()
    fig.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()


# ============ 2. Word文档生成 ============

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = __import__('docx.oxml', fromlist=['OxmlElement']).OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    """创建带格式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        # 背景色
        shading = __import__('docx.oxml', fromlist=['OxmlElement']).OxmlElement('w:shd')
        shading.set(qn('w:fill'), header_color)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if i % 2 == 1:
                shading = __import__('docx.oxml', fromlist=['OxmlElement']).OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'EBF5FB')
                shading.set(qn('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_heading_styled(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        rPr = run._r.get_or_add_rPr()
        rFonts = __import__('docx.oxml', fromlist=['OxmlElement']).OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '黑体')
        rPr.insert(0, rFonts)
    return heading


def add_para(doc, text, bold=False, indent=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '宋体'
    rPr = run._r.get_or_add_rPr()
    rFonts = __import__('docx.oxml', fromlist=['OxmlElement']).OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.insert(0, rFonts)
    if bold:
        run.font.bold = True
    return p


def add_image(doc, filepath, width_cm=14):
    """插入图片"""
    if os.path.exists(filepath):
        doc.add_picture(filepath, width=Cm(width_cm))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(doc, f"[图表占位：{os.path.basename(filepath)}，请在Word中插入对应图表]")


def generate_document():
    doc = Document()

    # ---- 页面设置 ----
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ---- 封面 ----
    for _ in range(4):
        doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("管理信息系统结课大作业")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.name = '黑体'
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("华中科技大学生活后勤系统升级改造设计方案\n——报修工单取消与预约改期子功能")
    run.font.size = Pt(16)
    run.font.name = '楷体'
    run.font.color.rgb = RGBColor(51, 51, 51)

    doc.add_paragraph()
    doc.add_paragraph()

    info_items = [
        f"姓    名：{STUDENT_NAME}",
        f"学    号：{STUDENT_ID}",
        "课程名称：管理信息系统",
        "选题系统：生活后勤系统（2244零星维修平台）",
        "功能聚焦：报修工单取消与预约改期子功能",
        "提交日期：2026年6月22日",
    ]
    for item in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        run.font.size = Pt(13)
        run.font.name = '宋体'

    doc.add_page_break()

    # ---- 目录占位 ----
    add_heading_styled(doc, "目录", level=1)
    add_para(doc, "（请在Word中插入自动目录：引用 → 目录 → 自动目录）")
    doc.add_page_break()

    # ==================== 第一部分：前期调研 ====================
    add_heading_styled(doc, "一、前期调研——现有系统分析", level=1)

    add_heading_styled(doc, "1.1 现有系统核心业务与功能梳理", level=2)
    add_para(doc, "2244零星维修平台是华中科技大学后勤服务的重要数字化入口，嵌套于企业微信「华中大微校园-生活后勤」板块中。学生需先选择校区（主校区/同济校区），方可进入维修报修主界面。基于现有系统截图，梳理其完整的业务链路与前端功能架构如下：", indent=True)

    add_para(doc, "（1）系统入口与导航结构", bold=True)
    add_para(doc, "系统整体采用微信内嵌H5形态，页面结构简单、层级较浅。首页仅包含「我要报修」与「服务指南」两个功能入口。底部导航栏设有「首页」「事项」「个人中心」三个Tab标签页。整体架构为轻量级工具型小程序，功能高度聚焦于报修提交。", indent=True)

    add_para(doc, "（2）现有业务链路", bold=True)
    add_para(doc, "现有系统仅支持「报修提交→工单存储」的正向单向流程，具体链路为：学生登录企业微信→进入生活后勤板块→选择校区→进入2244维修平台→点击「我要报修」→填写故障地址、故障项目、故障描述、上传图片→提交报修单→工单存入系统后台。提交后，学生可在「事项」页面查看个人报修单列表，但列表仅作信息展示，无任何工单操作按钮（无取消、无改期、无催单）。若需取消或修改报修，学生只能通过页面底部标注的固定维修电话进行线下沟通。", indent=True)

    add_para(doc, "（3）现有功能架构", bold=True)
    add_para(doc, "基于截图识别，现有系统前端功能模块如表1所示：", indent=True)

    add_styled_table(doc,
        ["功能模块", "页面/入口", "核心功能", "备注"],
        [
            ["报修提交", "首页→我要报修", "故障地址选择、故障项目选择、故障描述填写、图片上传、提交报修单", "仅支持正向提交"],
            ["报修单查看", "事项Tab页", "展示个人历史报修单列表、空状态提示", "无任何操作按钮"],
            [「服务指南」, "首页→服务指南", "展示《维修服务流程》《维修服务承诺》两份文件", "无取消/改期指引"],
            ["个人中心", "个人中心Tab页", "展示校园身份信息、「我的消息」入口", "无工单管理功能"],
            ["消息查看", "个人中心→我的消息", "支持关键词搜索历史消息", "无系统主动推送能力"],
            ["线下联系", "页面底部固定电话", "拨打维修电话进行线下沟通", "唯一的逆向操作渠道"],
        ],
        col_widths=[2.5, 3.0, 5.0, 3.0]
    )
    add_para(doc, "表1：现有2244零星维修平台前端功能模块一览", indent=True)

    add_para(doc, "综上，现有2244零星维修平台是一个轻量级、单向的报修工具，其核心业务能力仅覆盖「报修提交」「工单查看」两个环节。整个系统没有为学生提供任何工单级别的主动管理能力——无法在线上取消已提交的工单、无法修改预约的上门时间、无法查看工单实时状态变更。这一定位在维修需求简单、沟通成本低的场景下尚可运作，但当学生面临故障自行解决、临时不在宿舍、填报信息有误等实际情境时，系统的功能空白便转化为显著的体验痛点。", indent=True)

    add_heading_styled(doc, "1.2 学生体验痛点分析——聚焦取消与改期场景", level=2)
    add_para(doc, "站在学生用户体验视角，结合实际使用情境，还原三类典型痛点场景：", indent=True)

    add_para(doc, "场景一：故障自行解决，想要取消工单", bold=True)
    add_para(doc, "学生小A发现宿舍空调不制冷，通过2244平台提交了报修工单。提交后半小时，小A自行检查发现是遥控器电池没电，更换电池后空调恢复正常。小A试图取消已提交的工单，但在平台上找不到任何「取消」按钮。小A只能拨打页面底部维修电话，占线三次后才接通，口头说明情况后由对方手动记录取消。整个过程耗时约20分钟，且无任何书面确认凭证，无法确认工单是否真正被取消。", indent=True)

    add_para(doc, "场景二：临时有事不在宿舍，想改上门时间", bold=True)
    add_para(doc, "学生小B预约了周三下午14:00-16:00的维修上门服务，但临时接到导师通知需参加组会。小B想将上门时间改为周四上午，但在平台上找不到「修改时间」或「改期」入口。小B再次通过电话沟通，维修人员已出发前往宿舍，造成维修人员空跑一趟，小B也被要求在电话中反复解释原因，体验不佳。", indent=True)

    add_para(doc, "场景三：填报信息有误，想撤回重填", bold=True)
    add_para(doc, "学生小C在填写报修单时误选了错误的故障项目（将「水管漏水」选为「电路故障」），提交后才发现错误。由于系统无编辑或撤回功能，小C只能重新提交一份正确的工单，而错误的工单仍然存在于系统中，造成了工单冗余和维修资源的误分配。", indent=True)

    add_heading_styled(doc, "1.3 现有系统不足总结", level=2)
    add_para(doc, "基于上述三类典型场景的反推分析，现有2244零星维修平台在工单生命周期管理方面存在以下三点核心不足：", indent=True)

    add_para(doc, "不足一：缺少线上取消与改期操作入口，学生无工单自主管理权。现有系统仅提供「提交」这一种操作，学生在提交工单后便完全丧失了对工单的控制能力。取消和改期这两种最基本的工单管理需求，在系统中无任何线上入口，学生被迫转入线下电话渠道。这与当前高校信息化建设中「让数据多跑路、让师生少跑腿」的理念背道而驰。", indent=True)

    add_para(doc, "不足二：缺少标准化流程处理机制，线下电话沟通效率低且无留痕。现有的「电话取消/改期」模式完全依赖人工沟通，存在三大问题：电话可能占线或无人接听，沟通效率低；口头沟通无书面记录，后续追溯困难；取消/改期的处理结果取决于接听人员的个人判断，缺乏统一的业务规则（如工单已开工是否允许取消）。这种非标准化的处理方式既增加了后勤人员的工作负担，也无法保障学生的权益。", indent=True)

    add_para(doc, "不足三：缺少状态同步与消息通知机制，操作结果无法即时反馈。现有系统没有主动推送通知的能力，「我的消息」页面仅支持关键词搜索历史消息。学生提交取消/改期请求后，无法实时获知处理进度和结果：工单是否已被取消？改期是否已确认？维修人员是否已知晓变更？这些关键状态信息对学生而言完全不可见，形成了信息黑洞。", indent=True)

    add_heading_styled(doc, "1.4 改进方向", level=2)
    add_para(doc, "基于上述不足分析，本方案明确升级改造的范围边界：仅聚焦于在现有2244零星维修平台基础上，补充「报修工单取消与预约改期」这一单一子功能。该子功能不涉及报修提交优化、智能派单、服务评价等平台其他模块。具体改进方向包括：", indent=True)

    add_para(doc, "（1）补充线上操作入口：在「事项」页面的工单列表中，为每个处于可操作状态的工单增加「取消工单」与「改期预约」两个操作按钮，使学生能够在线上自主发起取消和改期请求。", indent=True)
    add_para(doc, "（2）建立流程化处理逻辑：设计从"学生发起申请→系统校验工单状态→自动/人工审核→工单状态更新→消息通知学生"的完整业务流程，明确各环节的处理规则（如工单已开工则不允许取消、改期时段需提前N小时等），确保每一笔取消/改期操作都有据可查。", indent=True)
    add_para(doc, "（3）实现状态同步与消息通知：通过企业微信消息推送能力，在工单状态发生变更时，实时通知报修学生和维修人员，确保信息同步、消除信息黑洞。同时在学生端的「事项」页面同步展示工单状态变更记录。", indent=True)

    doc.add_page_break()

    # ==================== 第二部分：新系统规划 ====================
    add_heading_styled(doc, "二、新系统规划——工单取消与改期子功能", level=1)

    add_heading_styled(doc, "2.1 组织模型", level=2)
    add_para(doc, "工单取消与预约改期子功能涉及三类核心角色：报修学生、后勤受理员、维修人员。各角色在取消/改期流程中的职责与协作关系如下：", indent=True)

    add_styled_table(doc,
        ["角色", "职责描述", "核心操作", "协作对象"],
        [
            ["报修学生", "报修工单的发起者和服务对象，在工单生命周期内拥有取消和改期的主动申请权", "发起取消申请、发起改期预约、查看申请进度、接收审核结果通知", "后勤受理员"],
            ["后勤受理员", "后勤管理部门的业务处理人员，负责审核学生的取消/改期申请，协调维修资源调度", "审核取消原因、审批改期时段、确认工单状态变更、手动处理异常工单", "报修学生、维修人员"],
            ["维修人员", "维修服务的执行者，在工单发生变更时接收通知并调整排班计划", "接收工单变更通知、确认新时段安排、反馈到场维修状态", "后勤受理员"],
        ],
        col_widths=[2.5, 4.0, 4.5, 2.5]
    )
    add_para(doc, "表2：组织角色职责与协作关系表", indent=True)

    # 插入组织模型图
    org_img = os.path.join(OUTPUT_DIR, "_fig_org_model.png")
    fig_organization_model(org_img)
    add_image(doc, org_img)

    add_heading_styled(doc, "2.2 C-U矩阵与子系统划分", level=2)
    add_para(doc, "围绕工单取消与预约改期功能，识别核心数据类与功能模块的创建（C）与使用（U）关系，构建C-U矩阵如下：", indent=True)

    add_styled_table(doc,
        ["数据类 \\ 功能模块", "取消申请提交", "改期申请提交", "申请审核处理", "工单状态更新", "消息推送"],
        [
            ["工单数据", "U", "U", "C/U", "C", "U"],
            ["预约时段数据", "", "U", "C/U", "", ""],
            ["取消原因数据", "U", "", "U", "", ""],
            ["消息数据", "", "", "", "", "C"],
        ],
        col_widths=[3.0, 2.0, 2.0, 2.0, 2.0, 2.0]
    )
    add_para(doc, "表3：工单取消与改期子功能C-U矩阵（C=创建 Create，U=使用 Use）", indent=True)

    add_para(doc, "基于C-U矩阵的聚类分析，将该子功能划分为三个子模块：", indent=True)

    add_styled_table(doc,
        ["子模块", "包含功能", "核心职责", "模块边界"],
        [
            ["学生端工单操作子模块", "取消申请提交、改期申请提交", "为学生提供线上取消和改期的操作入口，收集申请数据并提交至后台", "仅处理前端交互与数据收集，不涉及审核逻辑"],
            ["后台审核处理子模块", "申请审核处理、工单状态更新", "依据预设规则自动审核或转人工审核，执行工单状态变更，更新数据表", "仅处理后台业务逻辑，不直接面向学生"],
            ["消息通知子模块", "消息推送", "在工单状态变更后，向学生和维修人员推送企业微信消息通知", "仅负责消息触达，不参与业务决策"],
        ],
        col_widths=[3.0, 3.0, 4.0, 4.0]
    )
    add_para(doc, "表4：子系统划分与模块边界说明", indent=True)

    add_heading_styled(doc, "2.3 E-R图（实体-联系图）", level=2)
    add_para(doc, "识别工单取消与改期子功能涉及的五个核心实体：学生用户、报修工单、预约时段、取消原因、通知消息。各实体的核心属性及实体间联系如下：", indent=True)

    er_img = os.path.join(OUTPUT_DIR, "_fig_er.png")
    fig_er_diagram(er_img)
    add_image(doc, er_img)

    add_para(doc, "实体联系说明：", bold=True)
    add_para(doc, "（1）学生用户与报修工单：一对多（1:n）。一个学生可以提交多个报修工单，一个工单仅属于一个学生。", indent=True)
    add_para(doc, "（2）报修工单与预约时段：一对一（1:1）。一个工单对应一个预约上门时段，一个时段可被多个工单关联。", indent=True)
    add_para(doc, "（3）报修工单与取消原因：多对一（n:1）。多个工单的取消可关联同一个取消原因分类。", indent=True)
    add_para(doc, "（4）报修工单与通知消息：一对多（1:n）。一个工单的状态变更可触发多条通知消息。", indent=True)

    add_heading_styled(doc, "2.4 数据表设计", level=2)
    add_para(doc, "基于E-R图，设计以下四张核心数据表：", indent=True)

    add_para(doc, "（1）工单状态变更记录表（work_order_status_log）", bold=True)
    add_styled_table(doc,
        ["字段名", "数据类型", "主键/外键", "是否非空", "字段说明"],
        [
            ["log_id", "INT", "PK", "NOT NULL", "变更记录编号，自增主键"],
            ["order_id", "VARCHAR(20)", "FK", "NOT NULL", "工单编号，关联报修工单表"],
            ["student_id", "VARCHAR(15)", "FK", "NOT NULL", "学号，关联学生用户表"],
            ["from_status", "VARCHAR(2)", "", "NOT NULL", "变更前工单状态编码"],
            ["to_status", "VARCHAR(2)", "", "NOT NULL", "变更后工单状态编码"],
            ["change_type", "VARCHAR(10)", "", "NOT NULL", "变更类型：CANCEL取消/RESCHEDULE改期"],
            ["reason_code", "VARCHAR(4)", "FK", "", "取消/改期原因编码"],
            ["new_schedule_time", "DATETIME", "", "", "改期后新预约时间（取消时为空）"],
            ["operator_type", "VARCHAR(10)", "", "NOT NULL", "操作者类型：STUDENT/ADMIN"],
            ["change_time", "DATETIME", "", "NOT NULL", "状态变更时间"],
        ],
        col_widths=[2.8, 2.2, 2.0, 1.8, 5.0]
    )
    add_para(doc, "表5：工单状态变更记录表", indent=True)

    add_para(doc, "（2）预约时段配置表（schedule_slot_config）", bold=True)
    add_styled_table(doc,
        ["字段名", "数据类型", "主键/外键", "是否非空", "字段说明"],
        [
            ["slot_id", "INT", "PK", "NOT NULL", "时段编号，自增主键"],
            ["slot_date", "DATE", "", "NOT NULL", "预约日期"],
            ["start_time", "TIME", "", "NOT NULL", "时段开始时间，如14:00"],
            ["end_time", "TIME", "", "NOT NULL", "时段结束时间，如16:00"],
            ["max_capacity", "INT", "", "NOT NULL", "该时段最大可预约工单数"],
            ["current_count", "INT", "", "NOT NULL", "当前已预约工单数"],
            ["is_available", "TINYINT", "", "NOT NULL", "是否可选：1可选/0不可选"],
            ["campus", "VARCHAR(10)", "", "NOT NULL", "所属校区"],
        ],
        col_widths=[2.8, 2.2, 2.0, 1.8, 5.0]
    )
    add_para(doc, "表6：预约时段配置表", indent=True)

    add_para(doc, "（3）取消原因字典表（cancel_reason_dict）", bold=True)
    add_styled_table(doc,
        ["字段名", "数据类型", "主键/外键", "是否非空", "字段说明"],
        [
            ["reason_code", "VARCHAR(4)", "PK", "NOT NULL", "取消原因编码"],
            ["reason_desc", "VARCHAR(50)", "", "NOT NULL", "取消原因描述"],
            ["need_review", "TINYINT", "", "NOT NULL", "是否需要人工审核：1需要/0自动通过"],
            ["is_active", "TINYINT", "", "NOT NULL", "是否启用：1启用/0禁用"],
            ["sort_order", "INT", "", "", "显示排序号"],
        ],
        col_widths=[2.8, 2.2, 2.0, 1.8, 5.0]
    )
    add_para(doc, "表7：取消原因字典表", indent=True)

    add_para(doc, "（4）消息通知记录表（notification_record）", bold=True)
    add_styled_table(doc,
        ["字段名", "数据类型", "主键/外键", "是否非空", "字段说明"],
        [
            ["msg_id", "INT", "PK", "NOT NULL", "消息编号，自增主键"],
            ["order_id", "VARCHAR(20)", "FK", "NOT NULL", "关联工单编号"],
            ["student_id", "VARCHAR(15)", "FK", "NOT NULL", "接收消息的学生学号"],
            ["msg_type", "VARCHAR(20)", "", "NOT NULL", "消息类型：CANCEL_SUCCESS/RESCHEDULE_CONFIRM等"],
            ["msg_title", "VARCHAR(100)", "", "NOT NULL", "消息标题"],
            ["msg_content", "TEXT", "", "NOT NULL", "消息正文"],
            ["send_time", "DATETIME", "", "NOT NULL", "发送时间"],
            ["read_status", "TINYINT", "", "NOT NULL", "阅读状态：0未读/1已读"],
            ["read_time", "DATETIME", "", "", "阅读时间"],
            ["channel", "VARCHAR(20)", "", "NOT NULL", "推送渠道：WECOM企业微信/INAPP站内"],
        ],
        col_widths=[2.2, 2.0, 2.0, 1.8, 5.0]
    )
    add_para(doc, "表8：消息通知记录表", indent=True)

    doc.add_page_break()

    # ==================== 第三部分：新系统分析 ====================
    add_heading_styled(doc, "三、新系统分析——工单取消与改期子功能", level=1)

    add_heading_styled(doc, "3.1 业务流程图（泳道图）", level=2)
    add_para(doc, "采用泳道图展示「学生发起取消/改期申请→系统校验工单状态→自动/人工审核→工单状态更新→消息通知学生」的完整业务流程。业务涉及四个泳道：报修学生、系统平台、后勤受理员、维修人员。", indent=True)

    bf_img = os.path.join(OUTPUT_DIR, "_fig_biz_flow.png")
    fig_business_flow(bf_img)
    add_image(doc, bf_img)

    add_para(doc, "核心流程说明：", bold=True)
    add_para(doc, "（1）学生发起：学生在「事项」页面选择目标工单，点击「取消工单」或「改期预约」按钮，填写申请表单并提交。", indent=True)
    add_para(doc, "（2）系统校验：系统接收到申请后，首先校验工单当前状态。若工单状态为「已开工」（维修人员已出发），则直接拒绝操作，并提示学生通过电话紧急联系；若工单状态为「待受理」或"已受理（未开工）"，则进入审核环节。", indent=True)
    add_para(doc, "（3）自动/人工审核：对于「故障已自行解决」「信息填报错误」等简单取消原因，系统自动审核通过；对于涉及维修人员已排班等复杂改期场景，系统转人工审核，由后勤受理员确认。", indent=True)
    add_para(doc, "（4）状态更新与通知：审核通过后，系统更新工单状态，写入状态变更日志，并通过企业微信向学生推送审核结果通知，同时向维修人员推送工单变更通知。", indent=True)

    add_heading_styled(doc, "3.2 数据流程图（DFD）", level=2)
    add_para(doc, "采用分层数据流程图描述数据在系统中的流转逻辑：", indent=True)

    dfd0_img = os.path.join(OUTPUT_DIR, "_fig_dfd0.png")
    fig_data_flow_top(dfd0_img)
    add_image(doc, dfd0_img)

    dfd1_img = os.path.join(OUTPUT_DIR, "_fig_dfd1.png")
    fig_data_flow_level1(dfd1_img)
    add_image(doc, dfd1_img)

    add_para(doc, "数据流说明：", bold=True)
    add_para(doc, "（1）顶层DFD（DFD-0）：明确两个外部实体（报修学生、后勤受理员）与核心处理系统之间的数据流——学生向系统提交取消/改期申请并查询工单状态，系统向学生推送审核结果通知，后勤受理员接收变更通知。", indent=True)
    add_para(doc, "（2）一层DFD（DFD-1）：将核心处理拆解为三个子过程——P1申请提交处理（接收并校验学生申请）、P2审核处理（依据规则自动审核或转人工）、P3通知推送处理（生成并发送消息）。三个数据存储D1工单数据、D2取消原因字典、D3通知消息记录支撑整个数据流转。", indent=True)

    add_heading_styled(doc, "3.3 核心处理功能说明", level=2)
    add_para(doc, "（1）工单可取消/改期的状态阈值规则", bold=True)
    add_para(doc, "工单的生命周期状态包括：01-待受理、02-已受理（未派单）、03-已派单（未开工）、04-已开工（维修中）、05-已完成、06-已取消。其中，状态01-03范围内的工单允许学生发起取消或改期操作；状态04的工单不允许取消，仅可通过电话紧急联系；状态05和06的工单为终态，不再接受任何操作。该阈值规则的核心设计原则是：一旦维修人员已出发前往现场（状态04），取消将造成显著的资源浪费，学生端的自主操作权限应被收回。", indent=True)

    add_para(doc, "（2）改期时段的可选范围", bold=True)
    add_para(doc, "系统提供未来3个工作日内的可预约时段供学生选择，每个时段为2小时（如8:00-10:00、10:00-12:00、14:00-16:00、16:00-18:00）。每个时段设置最大容量上限（如每时段最多4单），当某时段已满时自动标记为不可选。改期申请提交时，系统校验目标时段的可用容量，容量不足时提示学生另选时段。该设计既保障了学生的选择自由度，也避免了维修资源的过度集中。", indent=True)

    add_para(doc, "（3）取消原因分类与审核规则", bold=True)
    add_styled_table(doc,
        ["取消原因编码", "取消原因描述", "审核方式", "说明"],
        [
            ["C01", "故障已自行解决", "自动通过", "学生自行修复，无需维修人员上门"],
            ["C02", "信息填报错误", "自动通过", "学生填写了错误的故障信息，允许直接取消"],
            ["C03", "临时不在宿舍", "自动通过", "学生因个人原因无法接待维修上门"],
            ["C04", "重复提交工单", "自动通过", "同一故障被重复提交"],
            ["C05", "其他原因", "人工审核", "不属于上述分类的特殊情况，转后勤受理员人工审核"],
        ],
        col_widths=[2.5, 3.5, 2.5, 5.0]
    )
    add_para(doc, "表9：取消原因分类与审核规则", indent=True)

    doc.add_page_break()

    # ==================== 第四部分：新系统设计 ====================
    add_heading_styled(doc, "四、新系统设计——工单取消与改期子功能", level=1)

    add_heading_styled(doc, "4.1 功能结构图", level=2)
    add_para(doc, "工单取消与改期子系统按「学生端-管理端」两大维度拆解，形成三个一级模块、九个二级功能、二十七个三级功能点的树状结构：", indent=True)

    fs_img = os.path.join(OUTPUT_DIR, "_fig_func_struct.png")
    fig_function_structure(fs_img)
    add_image(doc, fs_img)

    add_para(doc, "功能结构说明：", bold=True)
    add_para(doc, "学生端工单操作模块（一级）下设工单取消申请、预约改期申请、申请进度查询三个二级功能，每个二级功能再拆解为三个三级操作点，覆盖了从"发起申请→填写信息→查看结果"的完整学生操作链路。管理端审核处理模块（一级）下设取消原因审核、改期时段审批、工单状态管理三个二级功能。消息通知模块（一级）下设企业微信消息推送、站内通知管理、消息模板配置三个二级功能，承担整个子系统的消息触达职责。", indent=True)

    add_heading_styled(doc, "4.2 代码设计", level=2)
    add_para(doc, "针对工单取消与改期子功能涉及的三类核心编码，设计统一的编码规则：", indent=True)

    add_para(doc, "（1）工单状态编码", bold=True)
    add_styled_table(doc,
        ["编码", "状态名称", "说明", "允许的操作"],
        [
            ["01", "待受理", "工单已提交，尚未分派处理人员", "取消、改期"],
            ["02", "已受理", "工单已被受理，尚未指派维修人员", "取消、改期"],
            ["03", "已派单", "已指派维修人员，尚未开工", "取消、改期"],
            ["04", "已开工", "维修人员已出发/到达现场", "仅电话联系"],
            ["05", "已完成", "维修服务已确认完成", "不可操作（终态）"],
            ["06", "已取消", "工单已被取消", "不可操作（终态）"],
        ],
        col_widths=[1.5, 2.0, 5.0, 4.0]
    )
    add_para(doc, "表10：工单状态编码表", indent=True)

    add_para(doc, "（2）取消原因编码", bold=True)
    add_para(doc, "编码结构：C + 两位数字序号。如 C01=故障已自行解决，C02=信息填报错误，C03=临时不在宿舍，C04=重复提交工单，C05=其他原因。", indent=True)

    add_para(doc, "（3）改期申请单号编码", bold=True)
    add_para(doc, "编码结构：RS + 年月日(6位) + 四位流水号。共12位字符。示例：RS2606220001 表示2026年6月22日提交的第0001号改期申请。取消申请单号同理，前缀为CL。", indent=True)

    add_heading_styled(doc, "4.3 输入设计", level=2)
    add_para(doc, "设计两类核心输入表单，适配移动端H5操作场景：", indent=True)

    add_para(doc, "（1）工单取消申请表单", bold=True)
    add_styled_table(doc,
        ["序号", "输入项", "输入类型", "是否必填", "校验规则", "错误提示"],
        [
            ["1", "工单编号", "自动带入（只读）", "是", "系统自动填充当前工单编号", "—"],
            ["2", "工单摘要", "自动带入（只读）", "是", "显示故障项目+故障描述摘要", "—"],
            ["3", "取消原因", "下拉选择", "是", "从取消原因字典表取值", "请选择取消原因"],
            ["4", "补充说明", "多行文本", "否", "最多200字", "说明内容不超过200字"],
            ["5", "确认勾选", "复选框", "是", "用户需勾选「我已阅读取消须知」", "请确认取消须知"],
        ],
        col_widths=[1.0, 2.5, 3.0, 1.5, 3.0, 4.0]
    )
    add_para(doc, "表11：工单取消申请表单输入设计", indent=True)

    add_para(doc, "（2）改期预约申请表单", bold=True)
    add_styled_table(doc,
        ["序号", "输入项", "输入类型", "是否必填", "校验规则", "错误提示"],
        [
            ["1", "工单编号", "自动带入（只读）", "是", "系统自动填充", "—"],
            ["2", "原预约时间", "自动带入（只读）", "是", "显示当前预约的上门时间", "—"],
            ["3", "新预约日期", "日期选择器", "是", "可选范围：未来3个工作日内", "请选择有效的改期日期"],
            ["4", "新预约时段", "时段选择（按钮组）", "是", "根据日期加载可用时段，满额时段置灰", "请选择可用的时间段"],
            ["5", "改期原因", "下拉选择", "是", "从改期原因字典取值", "请选择改期原因"],
            ["6", "补充说明", "多行文本", "否", "最多200字", "说明内容不超过200字"],
        ],
        col_widths=[1.0, 2.5, 3.0, 1.5, 3.0, 4.0]
    )
    add_para(doc, "表12：改期预约申请表单输入设计", indent=True)

    add_heading_styled(doc, "4.4 输出设计", level=2)
    add_para(doc, "设计三类核心输出场景，明确输出内容、触发时机与展示载体：", indent=True)

    add_para(doc, "（1）学生端工单状态页输出", bold=True)
    add_styled_table(doc,
        ["输出项", "输出内容", "触发条件", "展示载体"],
        [
            ["工单状态标识", "当前工单状态：待受理/已受理/已派单/已开工/已取消等", "页面加载时查询", "事项页面→工单详情"],
            ["操作按钮", "根据状态阈值显示「取消工单」/「改期预约」按钮", "工单状态为01/02/03时显示", "事项页面→工单详情→操作区"],
            ["状态变更时间线", "按时间倒序展示工单的状态变更记录", "工单发生任何状态变更后", "事项页面→工单详情→变更记录"],
            ["申请进度提示", "显示「审核中」「已通过」「已拒绝」的进度状态", "学生提交申请后实时更新", "事项页面→工单详情→申请进度"],
        ],
        col_widths=[2.5, 5.0, 3.5, 3.5]
    )
    add_para(doc, "表13：学生端工单状态页输出设计", indent=True)

    add_para(doc, "（2）企业微信消息通知输出", bold=True)
    add_styled_table(doc,
        ["消息类型", "消息模板示例", "触发时机", "接收人"],
        [
            ["取消成功通知", "【报修取消】您提交的[空调维修]工单已成功取消。如有需要可重新报修。", "工单状态更新为「已取消」后", "报修学生"],
            ["改期确认通知", "【时间变更】您的[水管维修]工单上门时间已变更为6月23日14:00-16:00，请确保该时段有人在场。", "改期审核通过后", "报修学生"],
            ["改期拒绝通知", "【改期未通过】您的[电路维修]工单改期申请未通过，原因：所选时段已满，请重新选择。", "改期审核未通过时", "报修学生"],
            ["工单变更通知", "【工单变更】工单[WO20260622001]已被学生取消，原预约6月22日10:00-12:00时段已释放。", "工单取消/改期后", "维修人员"],
        ],
        col_widths=[2.5, 6.0, 3.0, 2.5]
    )
    add_para(doc, "表14：企业微信消息通知输出设计", indent=True)

    add_para(doc, "（3）管理端申请列表输出", bold=True)
    add_styled_table(doc,
        ["输出项", "输出内容", "触发条件", "展示载体"],
        [
            ["待审核列表", "按提交时间倒序排列的待审核取消/改期申请", "有新的取消/改期申请提交时", "后勤管理后台→申请审核页面"],
            ["申请详情", "工单信息+学生填写的取消/改期原因+补充说明", "管理员点击列表项时", "申请审核页面→申请详情弹窗"],
            ["审核操作区", 「提供」通过"「拒绝」按钮及审核意见输入框", "管理员查看申请详情时", "申请详情弹窗→底部操作区"],
            ["处理记录", "按时间倒序展示所有已处理的申请及审核结果", "页面加载时查询", "申请审核页面→历史记录Tab"],
        ],
        col_widths=[2.5, 5.0, 3.5, 3.5]
    )
    add_para(doc, "表15：管理端申请列表输出设计", indent=True)

    doc.add_page_break()

    # ==================== 结语 ====================
    add_heading_styled(doc, "五、总结与展望", level=1)
    add_para(doc, "本方案以华中科技大学2244零星维修平台为对象，聚焦「报修工单取消与预约改期」这一单一子功能，完成了从前期调研、系统规划、系统分析到系统设计的完整管理信息系统开发流程。方案的核心价值体现在以下三个方面：", indent=True)

    add_para(doc, "第一，补齐功能空白。现有2244平台仅支持正向报修提交，学生在工单生命周期内完全丧失自主管理权。本方案补充的取消与改期子功能，首次为学生提供了线上化的工单逆向操作能力，实现了从「提交后无能为力」到「全流程自主可控」的体验跃迁。", indent=True)

    add_para(doc, "第二，标准化业务流程。通过设计完整的泳道式业务流程、分层数据流程图和明确的状态阈值规则，将原本依赖线下电话的非标准化沟通，转变为系统化、可追溯、有据可查的线上流程。这不仅提升了学生的使用体验，也降低了后勤管理人员的沟通协调成本。", indent=True)

    add_para(doc, "第三，打通信息壁垒。通过企业微信消息推送能力，在工单状态发生变更时实时通知报修学生和维修人员，消除了原有「信息黑洞」问题。学生不再需要反复拨打电话确认进度，维修人员也能及时调整排班计划，实现供需双方的实时信息对称。", indent=True)

    add_para(doc, "方案的局限性与未来展望：本方案严格限定在工单取消与改期子功能范围内，未涉及智能派单、服务评价、维修质量追踪等平台其他模块。在后续版本迭代中，可考虑将取消/改期数据纳入后勤服务大数据分析，识别高频取消的故障类型和时段特征，为维修资源的优化配置提供数据支撑。此外，可探索引入智能客服机器人，进一步降低电话沟通渠道的依赖。", indent=True)

    doc.add_page_break()

    # ==================== 附录 ====================
    add_heading_styled(doc, "附录：图表索引", level=1)
    charts = [
        "表1：现有2244零星维修平台前端功能模块一览",
        "表2：组织角色职责与协作关系表",
        "表3：工单取消与改期子功能C-U矩阵",
        "表4：子系统划分与模块边界说明",
        "表5：工单状态变更记录表",
        "表6：预约时段配置表",
        "表7：取消原因字典表",
        "表8：消息通知记录表",
        "表9：取消原因分类与审核规则",
        "表10：工单状态编码表",
        "表11：工单取消申请表单输入设计",
        "表12：改期预约申请表单输入设计",
        "表13：学生端工单状态页输出设计",
        "表14：企业微信消息通知输出设计",
        "表15：管理端申请列表输出设计",
        "图1：组织角色关系模型",
        "图2：实体-联系图（E-R图）",
        "图3：泳道式业务流程图",
        "图4：顶层数据流程图（DFD-0）",
        "图5：一层数据流程图（DFD-1）",
        "图6：功能结构图（树状图）",
    ]
    for i, chart in enumerate(charts, 1):
        add_para(doc, f"{chart}")

    # ---- 保存 ----
    doc.save(OUTPUT_FILE)
    print(f"文档已生成：{OUTPUT_FILE}")
    return OUTPUT_FILE


if __name__ == "__main__":
    generate_document()
